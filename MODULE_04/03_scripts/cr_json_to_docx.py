#!/usr/bin/env python3
# cr_json_to_docx.py
# Génère un CR Chantier au format .docx à partir d'un JSON
# Deux modes: (1) avec template .docx, (2) génération programmatique

import json
import sys
import os
import glob
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_from_template(cr_data, template_path, output_path):
    """Génère un CR en remplissant un template .docx avec des marqueurs."""
    from docx_generator import generate_docx

    # Prépare les données de remplacement
    meta = cr_data["meta"]
    replacements = {
        "projet": meta["projet"],
        "date": meta["date"],
        "participants": ", ".join(meta["participants"]),
        "meteo": meta.get("meteo", "N/A"),
        "documents_consultes": ", ".join(meta.get("documents_consultes", [])),
        "taches_prevues": "; ".join(cr_data["avancement"]["taches_prevues"]),
        "taches_realisees": "; ".join(cr_data["avancement"]["taches_realisees"]),
        "ecarts": "; ".join(cr_data["avancement"]["ecarts"]) if cr_data["avancement"]["ecarts"] else "Aucun",
    }

    generate_docx(template_path, output_path, replacements)

def generate_programmatic(cr_data, output_path):
    """Génère un CR .docx de manière programmatique (sans template)."""
    doc = Document()
    meta = cr_data["meta"]

    # Titre principal
    title = doc.add_heading(f"Compte Rendu de Chantier", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Sous-titre avec projet et date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run(f"{meta['projet']} — {meta['date']}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 70, 127)

    doc.add_paragraph()  # Espace

    # Section Informations générales
    doc.add_heading("1. Informations générales", 1)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light List Accent 1'

    # Participants
    info_table.rows[0].cells[0].text = "Participants"
    info_table.rows[0].cells[1].text = ", ".join(meta["participants"])

    # Météo
    info_table.rows[1].cells[0].text = "Météo"
    info_table.rows[1].cells[1].text = meta.get("meteo", "N/A")

    # Documents consultés
    info_table.rows[2].cells[0].text = "Documents consultés"
    info_table.rows[2].cells[1].text = ", ".join(meta.get("documents_consultes", ["Aucun"]))

    # Rédacteur
    info_table.rows[3].cells[0].text = "Rédacteur"
    info_table.rows[3].cells[1].text = meta.get("redacteur", "N/A")

    # Met en gras la première colonne
    for row in info_table.rows:
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Section Avancement
    doc.add_heading("2. Avancement des travaux", 1)

    avancement = cr_data["avancement"]

    # Tâches prévues
    doc.add_heading("Tâches prévues", 2)
    for tache in avancement["taches_prevues"]:
        doc.add_paragraph(tache, style='List Bullet')

    # Tâches réalisées
    doc.add_heading("Tâches réalisées", 2)
    for tache in avancement["taches_realisees"]:
        p = doc.add_paragraph(tache, style='List Bullet')
        # Colore en vert
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 128, 0)

    # Écarts
    if avancement["ecarts"]:
        doc.add_heading("Écarts constatés", 2)
        for ecart in avancement["ecarts"]:
            p = doc.add_paragraph(ecart, style='List Bullet')
            # Colore en orange
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 140, 0)

    # Section Points remarquables
    if cr_data["points"]:
        doc.add_heading("3. Points remarquables", 1)

        # Groupe par type
        points_by_type = {}
        for point in cr_data["points"]:
            ptype = point["type"]
            if ptype not in points_by_type:
                points_by_type[ptype] = []
            points_by_type[ptype].append(point)

        for ptype, points in points_by_type.items():
            doc.add_heading(f"{ptype}", 2)

            for point in points:
                # Détermine la couleur selon la gravité
                color = RGBColor(0, 0, 0)
                if point["gravite"] == "critique":
                    color = RGBColor(255, 0, 0)
                elif point["gravite"] == "majeure":
                    color = RGBColor(255, 140, 0)
                elif point["gravite"] == "mineure":
                    color = RGBColor(255, 215, 0)

                p = doc.add_paragraph()
                run = p.add_run(f"[{point['gravite'].upper()}] {point['id']} — {point['description']}")
                run.font.color.rgb = color
                run.bold = True

                if point.get("liens"):
                    doc.add_paragraph(f"    Liens: {', '.join(point['liens'])}", style='List Bullet 2')

    # Section Photos
    if cr_data["photos"]:
        doc.add_heading("4. Photographies", 1)

        for photo in cr_data["photos"]:
            repere = f" ({photo['repere_plan']})" if photo.get("repere_plan") else ""
            commentaire = f" — {photo['commentaire']}" if photo.get("commentaire") else ""

            p = doc.add_paragraph()
            p.add_run(f"{photo['fichier']}{repere}{commentaire}")

            # Tente d'insérer l'image si elle existe
            if os.path.exists(photo['fichier']):
                try:
                    doc.add_picture(photo['fichier'], width=Inches(4))
                except Exception as e:
                    doc.add_paragraph(f"[Image non chargée: {e}]", style='Intense Quote')

    # Section Actions
    if cr_data["actions"]:
        doc.add_heading("5. Actions à mener", 1)

        # Tableau des actions
        action_table = doc.add_table(rows=1, cols=4)
        action_table.style = 'Light Grid Accent 1'

        # En-têtes
        headers = action_table.rows[0].cells
        headers[0].text = "Qui"
        headers[1].text = "Quoi"
        headers[2].text = "Quand"
        headers[3].text = "Critère de succès"

        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Lignes d'actions
        for action in cr_data["actions"]:
            row = action_table.add_row().cells
            row[0].text = action["qui"]
            row[1].text = action["quoi"]
            row[2].text = action["quand"]
            row[3].text = action.get("critere_succes", "—")

    # Pied de page
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer.add_run(f"Document généré le {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Sauvegarde
    doc.save(output_path)
    print(f"[OK] CR exporté -> {output_path}")

def find_default_template():
    """
    Cherche automatiquement un template .docx dans le répertoire 04_modeles/.
    Retourne le chemin du premier fichier .docx trouvé, ou None.
    """
    script_dir = os.path.dirname(os.path.abspath(__file__))
    modeles_dir = os.path.join(script_dir, "..", "04_modeles")

    # Normalise le chemin
    modeles_dir = os.path.normpath(modeles_dir)

    if not os.path.exists(modeles_dir):
        return None

    # Cherche tous les fichiers .docx
    pattern = os.path.join(modeles_dir, "*.docx")
    docx_files = glob.glob(pattern)

    # Filtre les fichiers temporaires Word (commence par ~$)
    docx_files = [f for f in docx_files if not os.path.basename(f).startswith('~$')]

    if docx_files:
        # Retourne le premier fichier trouvé
        return docx_files[0]

    return None

def main():
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python cr_json_to_docx.py cr.json out.docx")
        print("  python cr_json_to_docx.py cr.json out.docx --template template.docx")
        sys.exit(2)

    json_path = sys.argv[1]
    output_path = sys.argv[2]

    # Charge le JSON
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            cr_data = json.load(f)
    except Exception as e:
        print(f"[ERREUR] Impossible de lire {json_path}: {e}")
        sys.exit(1)

    # Détermine le template à utiliser
    template_path = None

    # Option 1 : --template fourni explicitement
    if "--template" in sys.argv:
        template_idx = sys.argv.index("--template") + 1
        if template_idx >= len(sys.argv):
            print("[ERREUR] Chemin du template manquant après --template")
            sys.exit(1)
        template_path = sys.argv[template_idx]
        print(f"[INFO] Utilisation du template spécifié: {template_path}")
    else:
        # Option 2 : cherche un template par défaut
        default_template = find_default_template()
        if default_template:
            template_path = default_template
            print(f"[INFO] Template détecté automatiquement: {os.path.basename(template_path)}")
            print(f"       Chemin complet: {template_path}")

    # Mode template ou programmatique?
    if template_path:
        # Vérifie que le template existe
        if not os.path.exists(template_path):
            print(f"[ERREUR] Template introuvable: {template_path}")
            sys.exit(1)
        generate_from_template(cr_data, template_path, output_path)
    else:
        print("[INFO] Aucun template trouvé, génération programmatique")
        generate_programmatic(cr_data, output_path)

if __name__ == "__main__":
    main()
