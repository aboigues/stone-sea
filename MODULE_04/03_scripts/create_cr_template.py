#!/usr/bin/env python3
# create_cr_template.py
# Crée un template .docx pour les CR Chantier avec des marqueurs {{variable}}

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_cr_template(output_path):
    """Crée un template .docx pour CR Chantier avec marqueurs."""
    doc = Document()

    # Titre principal
    title = doc.add_heading("Compte Rendu de Chantier", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Sous-titre avec marqueurs
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("{{projet}} — {{date}}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 70, 127)

    doc.add_paragraph()  # Espace

    # Section Informations générales
    doc.add_heading("1. Informations générales", 1)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light List Accent 1'

    # Participants
    info_table.rows[0].cells[0].text = "Participants"
    info_table.rows[0].cells[1].text = "{{participants}}"

    # Météo
    info_table.rows[1].cells[0].text = "Météo"
    info_table.rows[1].cells[1].text = "{{meteo}}"

    # Documents consultés
    info_table.rows[2].cells[0].text = "Documents consultés"
    info_table.rows[2].cells[1].text = "{{documents_consultes}}"

    # Rédacteur
    info_table.rows[3].cells[0].text = "Rédacteur"
    info_table.rows[3].cells[1].text = "{{redacteur}}"

    # Met en gras la première colonne
    for row in info_table.rows:
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Section Avancement
    doc.add_heading("2. Avancement des travaux", 1)

    doc.add_heading("Tâches prévues", 2)
    doc.add_paragraph("{{taches_prevues}}")

    doc.add_heading("Tâches réalisées", 2)
    doc.add_paragraph("{{taches_realisees}}")

    doc.add_heading("Écarts constatés", 2)
    doc.add_paragraph("{{ecarts}}")

    # Section Points remarquables
    doc.add_heading("3. Points remarquables", 1)
    doc.add_paragraph("{{points}}")
    doc.add_paragraph()
    doc.add_paragraph("Note: Pour des points multiples, utiliser la génération programmatique", style='Intense Quote')

    # Section Photos
    doc.add_heading("4. Photographies", 1)
    doc.add_paragraph("{{photos}}")

    # Section Actions
    doc.add_heading("5. Actions à mener", 1)
    doc.add_paragraph("{{actions}}")
    doc.add_paragraph()
    doc.add_paragraph("Note: Pour un tableau d'actions détaillé, utiliser la génération programmatique", style='Intense Quote')

    # Pied de page
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer.add_run("Document généré le {{date_generation}}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Instructions
    doc.add_page_break()
    doc.add_heading("Instructions d'utilisation du template", 0)

    instructions = [
        "Ce template utilise des marqueurs de type {{variable}} qui seront remplacés par les vraies valeurs.",
        "",
        "Marqueurs disponibles:",
        "  • {{projet}} - Nom du projet",
        "  • {{date}} - Date du CR",
        "  • {{participants}} - Liste des participants",
        "  • {{meteo}} - Conditions météo",
        "  • {{documents_consultes}} - Documents de référence",
        "  • {{redacteur}} - Nom du rédacteur",
        "  • {{taches_prevues}} - Tâches prévues",
        "  • {{taches_realisees}} - Tâches réalisées",
        "  • {{ecarts}} - Écarts constatés",
        "  • {{points}} - Points remarquables",
        "  • {{photos}} - Liste des photos",
        "  • {{actions}} - Actions à mener",
        "  • {{date_generation}} - Date de génération du document",
        "",
        "Usage avec docx_generator.py:",
        "  python docx_generator.py cr_template.docx output.docx data.json",
        "",
        "Usage avec cr_json_to_docx.py:",
        "  python cr_json_to_docx.py cr.json output.docx --template cr_template.docx",
        "",
        "Note: Pour des mises en forme complexes (tableaux, couleurs, images),",
        "préférer la génération programmatique sans template:",
        "  python cr_json_to_docx.py cr.json output.docx"
    ]

    for line in instructions:
        if line == "":
            doc.add_paragraph()
        elif line.startswith("  •"):
            doc.add_paragraph(line[4:], style='List Bullet')
        elif line.startswith("  python"):
            p = doc.add_paragraph(line)
            p.style = 'Intense Quote'
        else:
            doc.add_paragraph(line)

    # Sauvegarde
    doc.save(output_path)
    print(f"[OK] Template créé -> {output_path}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python create_cr_template.py output_template.docx")
        sys.exit(2)

    output_path = sys.argv[1]
    create_cr_template(output_path)

if __name__ == "__main__":
    main()
