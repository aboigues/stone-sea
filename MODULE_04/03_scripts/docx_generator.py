#!/usr/bin/env python3
# docx_generator.py
# Générateur générique de documents .docx à partir de templates
# Utilise des marqueurs {{variable}} dans le template .docx

import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def replace_in_paragraph(paragraph, replacements):
    """Remplace les marqueurs {{variable}} dans un paragraphe."""
    full_text = ''.join(run.text for run in paragraph.runs)

    # Cherche tous les marqueurs {{...}}
    for key, value in replacements.items():
        marker = "{{" + key + "}}"
        if marker in full_text:
            full_text = full_text.replace(marker, str(value))

    # Si modification, reconstitue le paragraphe
    if full_text != ''.join(run.text for run in paragraph.runs):
        # Garde le formatage du premier run
        if paragraph.runs:
            first_run = paragraph.runs[0]
            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            new_run.font.name = first_run.font.name
            new_run.font.size = first_run.font.size
            new_run.bold = first_run.bold
            new_run.italic = first_run.italic

def replace_in_table(table, replacements):
    """Remplace les marqueurs dans toutes les cellules d'un tableau."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)

def generate_docx(template_path, output_path, replacements):
    """
    Génère un document .docx à partir d'un template.

    Args:
        template_path: Chemin vers le template .docx
        output_path: Chemin du fichier de sortie
        replacements: Dict {marqueur: valeur} pour remplacer {{marqueur}}
    """
    doc = Document(template_path)

    # Remplace dans les paragraphes
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    # Remplace dans les tableaux
    for table in doc.tables:
        replace_in_table(table, replacements)

    # Sauvegarde
    doc.save(output_path)
    print(f"[OK] Document généré -> {output_path}")

def add_paragraph_to_doc(doc, text, style=None, bold=False, italic=False):
    """Ajoute un paragraphe formaté au document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if style:
        p.style = style
    return p

def add_heading_to_doc(doc, text, level=1):
    """Ajoute un titre au document."""
    return doc.add_heading(text, level=level)

def add_table_to_doc(doc, data, headers=None):
    """
    Ajoute un tableau au document.

    Args:
        doc: Document docx
        data: Liste de listes [[row1_col1, row1_col2], [row2_col1, row2_col2]]
        headers: Liste des en-têtes optionnelle
    """
    rows = len(data) + (1 if headers else 0)
    cols = len(data[0]) if data else len(headers) if headers else 0

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'

    # Headers
    if headers:
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(header)
            # Met en gras
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    # Data
    start_row = 1 if headers else 0
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.rows[start_row + i].cells[j].text = str(cell_data)

    return table

def main():
    if len(sys.argv) < 4:
        print("Usage: python docx_generator.py template.docx output.docx '{\"var1\":\"val1\", \"var2\":\"val2\"}'")
        print("  ou : python docx_generator.py template.docx output.docx data.json")
        sys.exit(2)

    template_path = sys.argv[1]
    output_path = sys.argv[2]
    data_input = sys.argv[3]

    # Parse les données (JSON string ou fichier)
    import json
    try:
        if data_input.startswith('{'):
            replacements = json.loads(data_input)
        else:
            with open(data_input, 'r', encoding='utf-8') as f:
                replacements = json.load(f)
    except json.JSONDecodeError as e:
        print(f"[ERREUR] Données JSON invalides: {e}")
        sys.exit(1)

    generate_docx(template_path, output_path, replacements)

if __name__ == "__main__":
    main()
